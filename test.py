import click
import requests
from bs4 import BeautifulSoup
from docx import Document
import pandas as pd
import numpy as np
from tqdm import tqdm

def search_eng(word: str):
    url = "https://endic.naver.com/search.nhn?sLn=kr&searchOption=all&query={}".format(word)
    res = requests.get(url)
    soup = BeautifulSoup(res.text, 'html.parser')

    try:
        word_box1 = soup.select('.word_num .list_e2')[0]
        meaning_ = word_box1.select('dd p .fnt_k05')[0].text  # 뜻
    except IndexError:
        print('"{}"에 대한 검색결과가 없습니다.'.format(word))
        return '"{}"에 대한 검색결과가 없습니다.', '', ''

    try:
        word_box2 = soup.select('.word_num .list_a')[0]
        ex_eng_ = word_box2.select('.fnt_e09._ttsText')[0].text  # 영어 예문
        ex_kor_ = word_box2.select('div.fnt_k10')[0].text  # 예문 해석
    except IndexError:
        return meaning_, '', ''

    return meaning_, ex_eng_, ex_kor_


if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--file', type=str, required=True, help='file to translate')
    args = parser.parse_args()

    filename = args.file
    document = Document(filename)
    table = document.tables[0]
    for idx in tqdm(range(1, len(table.rows)), desc='translation'):
        txt = table.cell(idx, 1).text
        meaning, ex_eng, ex_kor = search_eng(txt)
        trans = meaning + '\n' + ex_eng + ex_kor
        # print('word:', txt)
        # print(trans)
        # print('\n\n')
        table.cell(idx, 2).text = trans

    # print(len(document.tables[0].rows))
    save_name = '{}_trans.docx'.format(filename)
    document.save(save_name)





